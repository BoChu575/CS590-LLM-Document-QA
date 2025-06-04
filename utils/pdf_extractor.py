# utils/pdf_extractor.py
import PyPDF2
import docx
import logging
from io import BytesIO
from config import Config


class DocumentExtractor:
    """Document text extraction class supporting PDF, TXT, and DOCX files"""

    def __init__(self):
        """Initialize document extractor"""
        self.supported_types = Config.SUPPORTED_FILE_TYPES
        self.max_file_size = Config.MAX_FILE_SIZE

        # Set up logging
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

    def extract_text_from_file(self, file_path=None, file_content=None, file_type=None):
        """
        Extract text from various document types

        Args:
            file_path (str): Path to the file (for local files)
            file_content (bytes): File content in bytes (for uploaded files)
            file_type (str): File extension (pdf, txt, docx)

        Returns:
            tuple: (success: bool, text: str, error_message: str)
        """
        try:
            # Determine file type
            if file_path:
                file_type = file_path.split('.')[-1].lower()
                self.logger.info(f"Extracting text from file: {file_path}")
            elif file_content and file_type:
                file_type = file_type.lower()
                self.logger.info(f"Extracting text from uploaded {file_type} file")
            else:
                return False, "", "No file provided or file type not specified"

            # Check if file type is supported
            if file_type not in self.supported_types:
                return False, "", f"Unsupported file type: {file_type}. Supported types: {self.supported_types}"

            # Check file size if content is provided
            if file_content and len(file_content) > self.max_file_size:
                size_mb = len(file_content) / (1024 * 1024)
                return False, "", f"File too large: {size_mb:.2f}MB. Maximum size: {self.max_file_size / (1024 * 1024)}MB"

            # Extract text based on file type
            if file_type == 'pdf':
                return self._extract_pdf_text(file_path, file_content)
            elif file_type == 'txt':
                return self._extract_txt_text(file_path, file_content)
            elif file_type == 'docx':
                return self._extract_docx_text(file_path, file_content)
            else:
                return False, "", f"Extraction method not implemented for {file_type}"

        except Exception as e:
            error_msg = f"Error extracting text: {str(e)}"
            self.logger.error(error_msg)
            return False, "", error_msg

    def _extract_pdf_text(self, file_path=None, file_content=None):
        """Extract text from PDF file"""
        try:
            text = ""

            if file_path:
                # Read from file path
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        text += page_text + "\n"
                        self.logger.info(f"Extracted text from page {page_num + 1}")
            elif file_content:
                # Read from file content
                pdf_file = BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    self.logger.info(f"Extracted text from page {page_num + 1}")

            # Clean up text
            text = self._clean_text(text)

            if not text.strip():
                return False, "", "No text found in PDF file"

            self.logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return True, text, ""

        except Exception as e:
            return False, "", f"Error reading PDF: {str(e)}"

    def _extract_txt_text(self, file_path=None, file_content=None):
        """Extract text from TXT file"""
        try:
            if file_path:
                # Read from file path
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            elif file_content:
                # Read from file content
                try:
                    text = file_content.decode('utf-8')
                except UnicodeDecodeError:
                    # Try other encodings
                    for encoding in ['latin-1', 'ascii', 'cp1252']:
                        try:
                            text = file_content.decode(encoding)
                            break
                        except UnicodeDecodeError:
                            continue
                    else:
                        return False, "", "Unable to decode text file with common encodings"

            # Clean up text
            text = self._clean_text(text)

            if not text.strip():
                return False, "", "No text found in file"

            self.logger.info(f"Successfully extracted {len(text)} characters from TXT file")
            return True, text, ""

        except Exception as e:
            return False, "", f"Error reading TXT file: {str(e)}"

    def _extract_docx_text(self, file_path=None, file_content=None):
        """Extract text from DOCX file"""
        try:
            text = ""

            if file_path:
                # Read from file path
                doc = docx.Document(file_path)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
            elif file_content:
                # Read from file content
                docx_file = BytesIO(file_content)
                doc = docx.Document(docx_file)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"

            # Clean up text
            text = self._clean_text(text)

            if not text.strip():
                return False, "", "No text found in DOCX file"

            self.logger.info(f"Successfully extracted {len(text)} characters from DOCX file")
            return True, text, ""

        except Exception as e:
            return False, "", f"Error reading DOCX file: {str(e)}"

    def _clean_text(self, text):
        """Clean and normalize extracted text"""
        if not text:
            return ""

        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            # Strip whitespace and skip empty lines
            cleaned_line = line.strip()
            if cleaned_line:
                cleaned_lines.append(cleaned_line)

        # Join lines with single newlines
        cleaned_text = '\n'.join(cleaned_lines)

        # Remove excessive spaces
        import re
        cleaned_text = re.sub(r' +', ' ', cleaned_text)

        return cleaned_text

    def get_text_stats(self, text):
        """Get statistics about extracted text"""
        if not text:
            return {
                "character_count": 0,
                "word_count": 0,
                "line_count": 0,
                "paragraph_count": 0
            }

        # Calculate statistics
        character_count = len(text)
        word_count = len(text.split())
        line_count = len(text.split('\n'))
        paragraph_count = len([p for p in text.split('\n\n') if p.strip()])

        return {
            "character_count": character_count,
            "word_count": word_count,
            "line_count": line_count,
            "paragraph_count": paragraph_count
        }


# Test code
if __name__ == "__main__":
    print("🚀 Testing Document Extractor...")

    extractor = DocumentExtractor()

    # Test with sample text
    sample_text = """This is a sample document for testing.

    It contains multiple paragraphs with various content.
    This should help us test the text extraction functionality.

    The extractor should be able to handle different file types including PDF, TXT, and DOCX files.
    """

    # Test text cleaning
    cleaned = extractor._clean_text(sample_text)
    print(f"📝 Cleaned text length: {len(cleaned)} characters")

    # Test text statistics
    stats = extractor.get_text_stats(cleaned)
    print(f"📊 Text statistics: {stats}")

    print("✅ Document extractor ready for use!")
    print("📋 Supported file types:", extractor.supported_types)
    print(f"📏 Maximum file size: {extractor.max_file_size / (1024 * 1024):.1f}MB")